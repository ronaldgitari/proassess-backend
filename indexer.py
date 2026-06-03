"""
RAG Indexer — Stage 1 of the pipeline.

Responsibilities:
  1. Load documents from PDF, DOCX, XLSX, or web sources.
  2. Split into overlapping semantic chunks.
  3. Generate OpenAI embeddings.
  4. Upsert chunks into Chroma with rich metadata.
  5. Persist chunk provenance in PostgreSQL (DocumentChunk table).
"""

from __future__ import annotations

import hashlib
import io
import logging
import uuid
from typing import Any
from datetime import datetime

import httpx
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain.schema import Document

from config import settings

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────────────────────────
# Chroma client (singleton)
# ─────────────────────────────────────────────────────────────────

def get_chroma() -> Chroma:
    embeddings = OpenAIEmbeddings(
        model=settings.OPENAI_EMBEDDING_MODEL,
        dimensions=settings.OPENAI_EMBEDDING_DIMENSIONS,
        openai_api_key=settings.OPENAI_API_KEY,
    )
    return Chroma(
        collection_name=settings.CHROMA_COLLECTION,
        embedding_function=embeddings,
        collection_metadata={"hnsw:space": "cosine"},
    )


# ─────────────────────────────────────────────────────────────────
# Loaders
# ─────────────────────────────────────────────────────────────────

def load_pdf(file_bytes: bytes, source_name: str) -> list[Document]:
    """Extract text from a PDF, one Document per page."""
    import pypdf

    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    docs = []
    for page_num, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if not text:
            continue
        docs.append(Document(
            page_content=text,
            metadata={"source": source_name, "page": page_num, "type": "pdf"},
        ))
    logger.info("PDF '%s': extracted %d pages", source_name, len(docs))
    return docs


def load_docx(file_bytes: bytes, source_name: str) -> list[Document]:
    """Extract paragraphs from a Word document."""
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(file_bytes))
    full_text = []
    current_heading = "Document"

    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            current_heading = para.text.strip() or current_heading
        if para.text.strip():
            full_text.append((current_heading, para.text.strip()))

    # Group paragraphs by heading into logical sections
    sections: dict[str, list[str]] = {}
    for heading, text in full_text:
        sections.setdefault(heading, []).append(text)

    docs = []
    for heading, paragraphs in sections.items():
        content = "\n".join(paragraphs)
        docs.append(Document(
            page_content=content,
            metadata={"source": source_name, "section": heading, "type": "docx"},
        ))

    # Also extract tables
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        if rows:
            docs.append(Document(
                page_content="\n".join(rows),
                metadata={"source": source_name, "section": f"Table {i+1}", "type": "docx_table"},
            ))

    logger.info("DOCX '%s': extracted %d sections", source_name, len(docs))
    return docs


def load_xlsx(file_bytes: bytes, source_name: str) -> list[Document]:
    """
    Extract spreadsheet data sheet by sheet, treating each sheet as a
    document. Each row becomes a pipe-delimited text record.
    """
    import openpyxl

    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    docs = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        headers = []
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0:
                headers = [str(c) if c is not None else "" for c in row]
                continue
            if all(c is None for c in row):
                continue
            row_data = " | ".join(
                f"{headers[j] if j < len(headers) else j}: {str(v) if v is not None else ''}"
                for j, v in enumerate(row)
            )
            rows.append(row_data)

        if rows:
            docs.append(Document(
                page_content="\n".join(rows),
                metadata={"source": source_name, "sheet": sheet_name, "type": "xlsx"},
            ))

    logger.info("XLSX '%s': extracted %d sheets", source_name, len(docs))
    return docs


async def load_web(url: str, source_name: str) -> list[Document]:
    """
    Fetch a web page and extract main text content.
    For full-site crawling, extend this to follow internal links.
    """
    async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
        response = await client.get(url, headers={"User-Agent": "ProAssess-Indexer/1.0"})
        response.raise_for_status()

    soup = BeautifulSoup(response.text, "html.parser")

    # Remove non-content elements
    for tag in soup(["script", "style", "nav", "footer", "aside", "form"]):
        tag.decompose()

    # Prefer <article> or <main>, fall back to <body>
    main = soup.find("article") or soup.find("main") or soup.find("body")
    if not main:
        return []

    text = main.get_text(separator="\n")
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())

    logger.info("Web '%s': fetched %d chars", url, len(text))
    return [Document(
        page_content=text,
        metadata={"source": source_name, "url": url, "type": "web"},
    )]


# ─────────────────────────────────────────────────────────────────
# Chunker
# ─────────────────────────────────────────────────────────────────

def split_documents(docs: list[Document]) -> list[Document]:
    """Split documents into overlapping chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.CHUNK_SIZE,
        chunk_overlap=settings.CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(docs)
    logger.info("Split %d docs → %d chunks", len(docs), len(chunks))
    return chunks


# ─────────────────────────────────────────────────────────────────
# Main indexing function
# ─────────────────────────────────────────────────────────────────

async def index_source(
    *,
    source_id: str,
    org_id: str,
    domain_tag: str,
    docs: list[Document],
    db_session,          # AsyncSession — used to write DocumentChunk rows
) -> list[str]:
    """
    Chunk, embed, and store a list of pre-loaded Documents.

    Returns list of Chroma document IDs that were created.
    """
    from models.knowledge import DocumentChunk

    chunks = split_documents(docs)
    if not chunks:
        logger.warning("No chunks produced for source %s", source_id)
        return []

    # Annotate each chunk with org/source metadata for Chroma filtering
    for chunk in chunks:
        chunk.metadata.update({
            "org_id": str(org_id),
            "source_id": str(source_id),
            "domain_tag": domain_tag or "general",
        })

    # Build stable IDs: SHA1 of (source_id + chunk_index + first 100 chars)
    chroma_ids = []
    db_chunks = []
    for i, chunk in enumerate(chunks):
        raw = f"{source_id}_{i}_{chunk.page_content[:100]}"
        chroma_id = hashlib.sha1(raw.encode()).hexdigest()
        chroma_ids.append(chroma_id)

        db_chunks.append(DocumentChunk(
            id=uuid.uuid4(),
            source_id=uuid.UUID(source_id),
            chroma_id=chroma_id,
            content=chunk.page_content,
            metadata=chunk.metadata,
            token_count=len(chunk.page_content.split()),
        ))

    # Upsert to Chroma
    chroma = get_chroma()
    chroma.add_documents(documents=chunks, ids=chroma_ids)
    logger.info("Upserted %d chunks to Chroma for source %s", len(chunks), source_id)

    # Persist provenance to PostgreSQL
    db_session.add_all(db_chunks)
    await db_session.flush()

    return chroma_ids
